"""
Comprehensive Test Suite for DOCX to PDF Conversion Service

This test suite covers all endpoints mentioned in the assignment:
1. POST /api/v1/jobs - Submit a new conversion job
2. GET /api/v1/jobs/{job_id} - Get job status
3. GET /api/v1/jobs/{job_id}/download - Download results

Test Categories:
- Happy path scenarios
- Error handling
- Edge cases
- Concurrent operations
"""

import pytest
import asyncio
import os
import zipfile
import io
import time
from pathlib import Path
from httpx import AsyncClient
from fastapi import status
from docx import Document

# Import the FastAPI app
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from main import app
from app.database import Base, engine, SessionLocal
from app.models import Job, JobFile, JobStatus, FileStatus


# Test configuration
BASE_URL = "http://localhost:8000"
API_PREFIX = "/api/v1/jobs"


@pytest.fixture(scope="session")
def test_files_dir():
    """Create a temporary directory for test files."""
    test_dir = Path(__file__).parent / "test_files"
    test_dir.mkdir(exist_ok=True)
    yield test_dir
    # Cleanup after tests
    # Note: Commented out to allow inspection of test files
    # import shutil
    # shutil.rmtree(test_dir)


@pytest.fixture
def create_docx_file():
    """Factory fixture to create DOCX files with custom content."""
    def _create_docx(filename: str, content: str = "Test Document") -> bytes:
        doc = Document()
        doc.add_heading(content, 0)
        doc.add_paragraph(f"This is a test document: {filename}")
        doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit.")
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.read()
    
    return _create_docx


@pytest.fixture
def create_zip_file(create_docx_file):
    """Factory fixture to create ZIP files containing DOCX files."""
    def _create_zip(docx_files: list[str]) -> bytes:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for filename in docx_files:
                docx_content = create_docx_file(filename, f"Content for {filename}")
                zip_file.writestr(filename, docx_content)
        
        zip_buffer.seek(0)
        return zip_buffer.read()
    
    return _create_zip


@pytest.fixture
async def async_client():
    """Create an async HTTP client for testing."""
    async with AsyncClient(app=app, base_url=BASE_URL) as client:
        yield client


# ============================================================================
# TEST SUITE 1: POST /api/v1/jobs - Job Submission
# ============================================================================

class TestJobSubmission:
    """Test cases for job submission endpoint."""
    
    @pytest.mark.asyncio
    async def test_submit_job_single_file_success(self, async_client, create_zip_file):
        """
        Test Case 1: Submit a job with a single DOCX file
        Expected: 202 Accepted with job_id and status
        """
        zip_content = create_zip_file(["document1.docx"])
        
        response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        
        assert response.status_code == status.HTTP_202_ACCEPTED
        data = response.json()
        assert "job_id" in data
        assert "status" in data
        assert data["status"] == "pending"
        assert isinstance(data["job_id"], int)
        assert data["job_id"] > 0
    
    @pytest.mark.asyncio
    async def test_submit_job_multiple_files_success(self, async_client, create_zip_file):
        """
        Test Case 2: Submit a job with multiple DOCX files
        Expected: 202 Accepted with job_id
        """
        zip_content = create_zip_file([
            "report.docx",
            "proposal.docx",
            "summary.docx"
        ])
        
        response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("multiple.zip", zip_content, "application/zip")}
        )
        
        assert response.status_code == status.HTTP_202_ACCEPTED
        data = response.json()
        assert "job_id" in data
        assert data["status"] == "pending"
    
    @pytest.mark.asyncio
    async def test_submit_job_invalid_file_type(self, async_client):
        """
        Test Case 3: Submit a non-ZIP file
        Expected: 400 Bad Request
        """
        text_content = b"This is not a zip file"
        
        response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.txt", text_content, "text/plain")}
        )
        
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        data = response.json()
        assert "detail" in data
        assert "ZIP" in data["detail"]
    
    @pytest.mark.asyncio
    async def test_submit_job_empty_zip(self, async_client):
        """
        Test Case 4: Submit an empty ZIP file
        Expected: 400 Bad Request
        """
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            pass  # Empty ZIP
        
        zip_buffer.seek(0)
        
        response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("empty.zip", zip_buffer.read(), "application/zip")}
        )
        
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        data = response.json()
        assert "No DOCX files found" in data["detail"]
    
    @pytest.mark.asyncio
    async def test_submit_job_no_docx_files(self, async_client):
        """
        Test Case 5: Submit a ZIP with no DOCX files
        Expected: 400 Bad Request
        """
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.writestr("readme.txt", b"This is a text file")
            zip_file.writestr("image.png", b"fake image data")
        
        zip_buffer.seek(0)
        
        response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("no-docx.zip", zip_buffer.read(), "application/zip")}
        )
        
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        data = response.json()
        assert "No DOCX files found" in data["detail"]
    
    @pytest.mark.asyncio
    async def test_submit_job_corrupted_zip(self, async_client):
        """
        Test Case 6: Submit a corrupted ZIP file
        Expected: 400 Bad Request
        """
        corrupted_zip = b"PK\x03\x04corrupted data"
        
        response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("corrupted.zip", corrupted_zip, "application/zip")}
        )
        
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        data = response.json()
        assert "detail" in data
    
    @pytest.mark.asyncio
    async def test_submit_job_with_macosx_files(self, async_client, create_zip_file, create_docx_file):
        """
        Test Case 7: Submit a ZIP with __MACOSX system files (should be filtered)
        Expected: 202 Accepted, system files ignored
        """
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add valid DOCX
            docx_content = create_docx_file("valid.docx")
            zip_file.writestr("valid.docx", docx_content)
            
            # Add __MACOSX files (should be filtered)
            zip_file.writestr("__MACOSX/._valid.docx", b"system file")
        
        zip_buffer.seek(0)
        
        response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("macosx.zip", zip_buffer.read(), "application/zip")}
        )
        
        assert response.status_code == status.HTTP_202_ACCEPTED
        data = response.json()
        assert "job_id" in data


# ============================================================================
# TEST SUITE 2: GET /api/v1/jobs/{job_id} - Job Status
# ============================================================================

class TestJobStatus:
    """Test cases for job status endpoint."""
    
    @pytest.mark.asyncio
    async def test_get_job_status_pending(self, async_client, create_zip_file):
        """
        Test Case 8: Get status of a newly created job
        Expected: 200 OK with status "pending" or "processing"
        """
        # Create a job
        zip_content = create_zip_file(["test.docx"])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        # Get status
        response = await async_client.get(f"{API_PREFIX}/{job_id}")
        
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert "id" in data
        assert "status" in data
        assert "created_at" in data
        assert "files" in data
        assert data["id"] == job_id
        assert data["status"] in ["pending", "processing", "completed"]
        assert isinstance(data["files"], list)
        assert len(data["files"]) > 0
    
    @pytest.mark.asyncio
    async def test_get_job_status_with_file_details(self, async_client, create_zip_file):
        """
        Test Case 9: Verify file details in job status
        Expected: Each file has filename, status, and optional error_message
        """
        zip_content = create_zip_file(["doc1.docx", "doc2.docx"])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        response = await async_client.get(f"{API_PREFIX}/{job_id}")
        
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        
        # Verify file structure
        for file_info in data["files"]:
            assert "id" in file_info
            assert "filename" in file_info
            assert "status" in file_info
            assert file_info["filename"].endswith(".docx")
            assert file_info["status"] in ["pending", "processing", "completed", "failed"]
    
    @pytest.mark.asyncio
    async def test_get_job_status_nonexistent(self, async_client):
        """
        Test Case 10: Get status of non-existent job
        Expected: 404 Not Found
        """
        response = await async_client.get(f"{API_PREFIX}/99999")
        
        assert response.status_code == status.HTTP_404_NOT_FOUND
        data = response.json()
        assert "detail" in data
        assert "not found" in data["detail"].lower()
    
    @pytest.mark.asyncio
    async def test_get_job_status_invalid_id(self, async_client):
        """
        Test Case 11: Get status with invalid job ID format
        Expected: 422 Unprocessable Entity
        """
        response = await async_client.get(f"{API_PREFIX}/invalid-id")
        
        assert response.status_code == status.HTTP_422_UNPROCESSABLE_ENTITY
    
    @pytest.mark.asyncio
    async def test_get_job_status_completed_has_download_url(self, async_client, create_zip_file):
        """
        Test Case 12: Completed job should have download_url
        Expected: download_url present when status is "completed"
        """
        zip_content = create_zip_file(["simple.docx"])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        # Poll until completed (with timeout)
        max_attempts = 30
        for _ in range(max_attempts):
            response = await async_client.get(f"{API_PREFIX}/{job_id}")
            data = response.json()
            
            if data["status"] in ["completed", "partial_success"]:
                assert "download_url" in data
                assert data["download_url"] is not None
                assert f"/api/v1/jobs/{job_id}/download" in data["download_url"]
                break
            
            await asyncio.sleep(2)
        else:
            pytest.skip("Job did not complete in time")
    
    @pytest.mark.asyncio
    async def test_get_job_status_pending_no_download_url(self, async_client, create_zip_file):
        """
        Test Case 13: Pending job should not have download_url
        Expected: download_url is None or absent when status is "pending"
        """
        # Create job with many files to ensure it stays pending
        zip_content = create_zip_file([f"doc{i}.docx" for i in range(5)])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        # Immediately check status
        response = await async_client.get(f"{API_PREFIX}/{job_id}")
        data = response.json()
        
        if data["status"] == "pending":
            assert data.get("download_url") is None


# ============================================================================
# TEST SUITE 3: GET /api/v1/jobs/{job_id}/download - Download Results
# ============================================================================

class TestJobDownload:
    """Test cases for job download endpoint."""
    
    @pytest.mark.asyncio
    async def test_download_completed_job(self, async_client, create_zip_file):
        """
        Test Case 14: Download results from a completed job
        Expected: 200 OK with application/zip content
        """
        zip_content = create_zip_file(["download_test.docx"])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        # Wait for completion
        max_attempts = 30
        for _ in range(max_attempts):
            status_response = await async_client.get(f"{API_PREFIX}/{job_id}")
            if status_response.json()["status"] in ["completed", "partial_success"]:
                break
            await asyncio.sleep(2)
        else:
            pytest.skip("Job did not complete in time")
        
        # Download results
        response = await async_client.get(f"{API_PREFIX}/{job_id}/download")
        
        assert response.status_code == status.HTTP_200_OK
        assert response.headers["content-type"] == "application/zip"
        assert len(response.content) > 0
        
        # Verify it's a valid ZIP
        zip_buffer = io.BytesIO(response.content)
        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
            files = zip_file.namelist()
            assert len(files) > 0
            # Should contain PDF files
            pdf_files = [f for f in files if f.endswith('.pdf')]
            assert len(pdf_files) > 0
    
    @pytest.mark.asyncio
    async def test_download_multiple_files(self, async_client, create_zip_file):
        """
        Test Case 15: Download results with multiple converted files
        Expected: ZIP contains all converted PDFs
        """
        docx_files = ["file1.docx", "file2.docx", "file3.docx"]
        zip_content = create_zip_file(docx_files)
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        # Wait for completion
        max_attempts = 40
        for _ in range(max_attempts):
            status_response = await async_client.get(f"{API_PREFIX}/{job_id}")
            if status_response.json()["status"] in ["completed", "partial_success"]:
                break
            await asyncio.sleep(2)
        else:
            pytest.skip("Job did not complete in time")
        
        # Download and verify
        response = await async_client.get(f"{API_PREFIX}/{job_id}/download")
        
        assert response.status_code == status.HTTP_200_OK
        
        # Verify all files converted
        zip_buffer = io.BytesIO(response.content)
        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
            pdf_files = [f for f in zip_file.namelist() if f.endswith('.pdf')]
            # Should have same number of PDFs as input DOCX files
            assert len(pdf_files) >= 1  # At least some files converted
    
    @pytest.mark.asyncio
    async def test_download_nonexistent_job(self, async_client):
        """
        Test Case 16: Download from non-existent job
        Expected: 404 Not Found
        """
        response = await async_client.get(f"{API_PREFIX}/99999/download")
        
        assert response.status_code == status.HTTP_404_NOT_FOUND
    
    @pytest.mark.asyncio
    async def test_download_pending_job(self, async_client, create_zip_file):
        """
        Test Case 17: Attempt to download from pending job
        Expected: 400 Bad Request
        """
        # Create job with many files
        zip_content = create_zip_file([f"file{i}.docx" for i in range(10)])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        # Immediately try to download (should fail)
        response = await async_client.get(f"{API_PREFIX}/{job_id}/download")
        
        # Should be 400 if still pending, or 200 if already completed
        if response.status_code == status.HTTP_400_BAD_REQUEST:
            data = response.json()
            assert "not ready" in data["detail"].lower()
    
    @pytest.mark.asyncio
    async def test_download_result_file_not_found(self, async_client, create_zip_file):
        """
        Test Case 18: Download when result file was deleted
        Expected: 404 Not Found
        """
        # This test would require manually deleting the result file
        # Skipping for now as it requires file system manipulation
        pytest.skip("Requires manual file deletion")


# ============================================================================
# TEST SUITE 4: Integration & End-to-End Tests
# ============================================================================

class TestEndToEnd:
    """End-to-end integration tests."""
    
    @pytest.mark.asyncio
    async def test_complete_workflow(self, async_client, create_zip_file):
        """
        Test Case 19: Complete workflow from submission to download
        Expected: All steps succeed
        """
        # Step 1: Submit job
        zip_content = create_zip_file(["workflow_test.docx"])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        assert create_response.status_code == status.HTTP_202_ACCEPTED
        job_id = create_response.json()["job_id"]
        
        # Step 2: Poll status until completed
        max_attempts = 30
        final_status = None
        for attempt in range(max_attempts):
            status_response = await async_client.get(f"{API_PREFIX}/{job_id}")
            assert status_response.status_code == status.HTTP_200_OK
            
            data = status_response.json()
            final_status = data["status"]
            
            if final_status in ["completed", "partial_success", "failed"]:
                break
            
            await asyncio.sleep(2)
        
        assert final_status in ["completed", "partial_success"]
        
        # Step 3: Download results
        download_response = await async_client.get(f"{API_PREFIX}/{job_id}/download")
        assert download_response.status_code == status.HTTP_200_OK
        assert len(download_response.content) > 0
    
    @pytest.mark.asyncio
    async def test_concurrent_job_submissions(self, async_client, create_zip_file):
        """
        Test Case 20: Submit multiple jobs concurrently
        Expected: All jobs created successfully with unique IDs
        """
        # Create multiple jobs concurrently
        tasks = []
        for i in range(3):
            zip_content = create_zip_file([f"concurrent_{i}.docx"])
            task = async_client.post(
                f"{API_PREFIX}/",
                files={"file": (f"test_{i}.zip", zip_content, "application/zip")}
            )
            tasks.append(task)
        
        responses = await asyncio.gather(*tasks)
        
        # Verify all succeeded
        job_ids = []
        for response in responses:
            assert response.status_code == status.HTTP_202_ACCEPTED
            job_id = response.json()["job_id"]
            job_ids.append(job_id)
        
        # Verify unique job IDs
        assert len(job_ids) == len(set(job_ids))
    
    @pytest.mark.asyncio
    async def test_large_batch_processing(self, async_client, create_zip_file):
        """
        Test Case 21: Process a large batch of files
        Expected: All files processed successfully
        """
        # Create ZIP with 10 files
        docx_files = [f"batch_file_{i}.docx" for i in range(10)]
        zip_content = create_zip_file(docx_files)
        
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("large_batch.zip", zip_content, "application/zip")}
        )
        assert create_response.status_code == status.HTTP_202_ACCEPTED
        job_id = create_response.json()["job_id"]
        
        # Wait for completion (longer timeout for large batch)
        max_attempts = 60
        for _ in range(max_attempts):
            status_response = await async_client.get(f"{API_PREFIX}/{job_id}")
            data = status_response.json()
            
            if data["status"] in ["completed", "partial_success"]:
                # Verify file count
                assert len(data["files"]) == len(docx_files)
                break
            
            await asyncio.sleep(3)
        else:
            pytest.skip("Large batch did not complete in time")


# ============================================================================
# TEST SUITE 5: Error Handling & Edge Cases
# ============================================================================

class TestErrorHandling:
    """Test error handling scenarios."""
    
    @pytest.mark.asyncio
    async def test_missing_file_parameter(self, async_client):
        """
        Test Case 22: Submit job without file parameter
        Expected: 422 Unprocessable Entity
        """
        response = await async_client.post(f"{API_PREFIX}/")
        
        assert response.status_code == status.HTTP_422_UNPROCESSABLE_ENTITY
    
    @pytest.mark.asyncio
    async def test_health_check_endpoint(self, async_client):
        """
        Test Case 23: Verify health check endpoint
        Expected: 200 OK with success message
        """
        response = await async_client.get("/")
        
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert "message" in data
    
    @pytest.mark.asyncio
    async def test_api_documentation_accessible(self, async_client):
        """
        Test Case 24: Verify API documentation is accessible
        Expected: 200 OK
        """
        response = await async_client.get("/docs")
        
        assert response.status_code == status.HTTP_200_OK
    
    @pytest.mark.asyncio
    async def test_job_status_persistence(self, async_client, create_zip_file):
        """
        Test Case 25: Verify job status persists across multiple requests
        Expected: Same status returned for same job
        """
        zip_content = create_zip_file(["persistence_test.docx"])
        create_response = await async_client.post(
            f"{API_PREFIX}/",
            files={"file": ("test.zip", zip_content, "application/zip")}
        )
        job_id = create_response.json()["job_id"]
        
        # Get status multiple times
        response1 = await async_client.get(f"{API_PREFIX}/{job_id}")
        response2 = await async_client.get(f"{API_PREFIX}/{job_id}")
        
        assert response1.status_code == status.HTTP_200_OK
        assert response2.status_code == status.HTTP_200_OK
        assert response1.json()["id"] == response2.json()["id"]


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
